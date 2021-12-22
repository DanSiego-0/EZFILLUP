import docx

doc = docx.Document('text.docx')
substring = "_____"
replacement_word = None


def GetNumPar(array):
    for items in range(len(array)):
        items+=1
    if items > 0:
        items-=1
    if items < 0:
        items -= 1
    return items

def GetRuns(x):
    for run in range(len(doc.paragraphs[x].runs)):
        run+=1
    return run

print(f'Number of Paragraphs: {GetNumPar(doc.paragraphs)} \n Number of Runs: {GetRuns(GetNumPar(doc.paragraphs))}')
numofp = GetNumPar(doc.paragraphs)
numofr = None
if numofp <=1:
    numofp+=1
elif numofp > 1:
    numofp+=1

for par in range(0,numofp):
    numofr = GetRuns(par)
    for r in range(0,numofr):
        if doc.paragraphs[par].runs[r].text != None and substring in doc.paragraphs[par].runs[r].text:
            print(f"{par}'s runs:{doc.paragraphs[par].runs[r-1].text} {doc.paragraphs[par].runs[r].text}")
            doc.paragraphs[par].runs[r].text = input("Write here: ") + " "
            
doc.save('Updated.docx')


